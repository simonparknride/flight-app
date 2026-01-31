import streamlit as st
import re, io
from datetime import datetime, timedelta, time as dtime
from typing import List, Dict, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm

# ---------------- Page config ----------------
st.set_page_config(page_title="Flight List Factory", layout="centered")

# ---------------- Regex / constants ----------------
TIME_LINE = re.compile(r"^(\d{1,2}:\d{2}\s?[AP]M)\s+([A-Z0-9]{2,4}\d*)", re.I)
DATE_HEADER = re.compile(r"^[A-Za-z]+,\s+\w+\s+\d{1,2}$")
IATA_IN_PARENS = re.compile(r"\(([^)]+)\)")
REGO_LIKE = re.compile(r"^[A-Za-z0-9\-–—]+$")

ALLOWED_AIRLINES = {"NZ","QF","JQ","CZ","CA","SQ","LA","IE","FX"}
NZ_DOMESTIC = {"AKL","WLG","CHC","ZQN","TRG","NPE","PMR","NSN","NPL","DUD","IVC","TUO","WRE","BHE","ROT","GIS","KKE","WHK","WAG","PPQ"}

# ---------------- Parsing ----------------
def try_parse_date(line: str, year: int):
    for fmt in ("%A, %b %d %Y", "%A, %B %d %Y", "%a, %b %d %Y"):
        try:
            return datetime.strptime(f"{line} {year}", fmt).date()
        except:
            pass
    return None

def parse_raw_lines(lines: List[str], year: int):
    records = []
    current_date = None
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if DATE_HEADER.match(line):
            current_date = try_parse_date(line, year)
            i += 1
            continue

        m = TIME_LINE.match(line)
        if m and current_date and i + 2 < len(lines):
            time_str, flight = m.groups()
            dest_line = lines[i+1]
            extra_line = lines[i+2]

            dest = ""
            m2 = IATA_IN_PARENS.search(dest_line)
            if m2:
                dest = m2.group(1).upper()

            reg = ""
            for p in reversed(IATA_IN_PARENS.findall(extra_line)):
                if REGO_LIKE.match(p):
                    reg = p
                    break

            try:
                t = datetime.strptime(time_str.replace(" ", ""), "%I:%M%p").time()
                dt = datetime.combine(current_date, t)
            except:
                dt = None

            records.append({
                "dt": dt,
                "time": time_str,
                "flight": flight,
                "dest": dest,
                "reg": reg
            })
            i += 3
            continue

        i += 1

    # -------- 연말/연초 보정 --------
    for idx in range(1, len(records)):
        if records[idx]["dt"] and records[idx-1]["dt"]:
            if records[idx]["dt"] < records[idx-1]["dt"]:
                records[idx]["dt"] += timedelta(days=1)

    return records

# ---------------- Filtering ----------------
def filter_records(records, start_t, end_t):
    if not records:
        return [], None, None

    base_date = records[0]["dt"].date()
    start_dt = datetime.combine(base_date, start_t)
    end_dt = datetime.combine(
        base_date + (timedelta(days=1) if end_t <= start_t else timedelta()),
        end_t
    )

    out = []
    for r in records:
        if not r["dt"]:
            continue
        if r["flight"][:2] not in ALLOWED_AIRLINES:
            continue
        if r["dest"] in NZ_DOMESTIC:
            continue
        if start_dt <= r["dt"] <= end_dt:
            out.append(r)

    return out, start_dt, end_dt

# ---------------- DOCX ----------------
def build_docx(records, start_dt, end_dt):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{start_dt:%d %b} – {end_dt:%d %b}").bold = True

    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r in records:
        row = table.add_row().cells
        row[0].text = r["flight"]
        row[1].text = r["dt"].strftime("%H:%M")
        row[2].text = r["dest"]
        row[3].text = r["reg"]

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ---------------- PDF ----------------
def build_pdf(records):
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    y = 800

    for r in records:
        c.drawString(50, y, f'{r["flight"]} {r["dt"].strftime("%H:%M")} {r["dest"]} {r["reg"]}')
        y -= 20
        if y < 50:
            c.showPage()
            y = 800

    c.save()
    buf.seek(0)
    return buf

# ---------------- UI ----------------
st.title("Flight List Factory")

year = st.number_input("Year", value=datetime.now().year)
s_time = st.time_input("Start time", dtime(5,0))
e_time = st.time_input("End time", dtime(4,55))

file = st.file_uploader("Upload raw txt", type="txt")

if file:
    lines = file.read().decode("utf-8", "ignore").splitlines()
    recs = parse_raw_lines(lines, year)
    filt, sdt, edt = filter_records(recs, s_time, e_time)

    if filt:
        st.success(f"{len(filt)} flights")
        st.download_button("Download DOCX", build_docx(filt, sdt, edt), "flights.docx")
        st.download_button("Download PDF", build_pdf(filt), "labels.pdf")
        st.table([
            {
                "Flight": r["flight"],
                "Time": r["dt"].strftime("%H:%M"),
                "Dest": r["dest"],
                "Reg": r["reg"]
            } for r in filt
        ])
    else:
        st.warning("No matching flights")
